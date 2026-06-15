"""Export LLD markdown to .docx (Word)."""
import io
import re
from docx import Document
from docx.shared import Pt, RGBColor


def markdown_to_docx(title: str, markdown: str) -> bytes:
    doc = Document()

    # Default style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    doc.add_heading(title, level=0)

    lines = markdown.split("\n")
    in_code = False
    code_buf: list[str] = []

    def flush_code():
        nonlocal code_buf
        if code_buf:
            p = doc.add_paragraph()
            run = p.add_run("\n".join(code_buf))
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            code_buf = []

    for raw_line in lines:
        line = raw_line.rstrip("\n")

        # Code blocks
        if line.strip().startswith("```"):
            if in_code:
                flush_code()
                in_code = False
            else:
                in_code = True
            continue
        if in_code:
            code_buf.append(line)
            continue

        # Strip HTML anchors / tags
        clean = re.sub(r"<[^>]+>", "", line).rstrip()

        if not clean.strip():
            continue

        # Headings
        m = re.match(r"^(#{1,6})\s+(.*)$", clean)
        if m:
            level = min(len(m.group(1)), 4)
            doc.add_heading(m.group(2).strip(), level=level)
            continue

        # Tables (simple pipe table)
        if clean.lstrip().startswith("|") and "|" in clean[1:]:
            # Not implementing full table parsing — render as monospace block
            p = doc.add_paragraph()
            run = p.add_run(clean)
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            continue

        # Lists
        list_match = re.match(r"^(\s*)([-*+]|\d+\.)\s+(.*)$", clean)
        if list_match:
            content = list_match.group(3)
            doc.add_paragraph(_strip_inline_md(content), style="List Bullet")
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        _add_inline(p, clean)

    flush_code()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _strip_inline_md(text: str) -> str:
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    return text


def _add_inline(paragraph, text: str) -> None:
    # Very small inline markdown handler: **bold**, *italic*, `code`
    tokens = re.split(r"(\*\*.+?\*\*|\*.+?\*|`.+?`)", text)
    for tok in tokens:
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            run = paragraph.add_run(tok[2:-2])
            run.bold = True
        elif tok.startswith("*") and tok.endswith("*"):
            run = paragraph.add_run(tok[1:-1])
            run.italic = True
        elif tok.startswith("`") and tok.endswith("`"):
            run = paragraph.add_run(tok[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        else:
            paragraph.add_run(tok)
